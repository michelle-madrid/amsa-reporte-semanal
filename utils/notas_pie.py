"""Traspaso de notas al pie desde los Word fuente al Word final.

python-docx no soporta notas al pie de forma nativa, pero la plantilla
(Template Viñetas Python.docx) ya trae la parte word/footnotes.xml, los estilos
('Refdenotaalpie' para la referencia, 'Textonotapie' para el texto) y la relación.
Aquí se: (1) leen las notas del Word fuente con su texto de anclaje, y (2) se
reinsertan como notas al pie reales en el Word final, ubicándolas tras el mismo
texto donde estaban en el origen.
"""
import copy

from lxml import etree
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_w = "{%s}" % _W

# Tipos de "footnote" que no son notas reales (separadores/avisos de Word).
_TIPOS_NO_NOTA = {"separator", "continuationSeparator", "continuationNotice"}

# Estilos de la plantilla (Word en español).
_ESTILO_REF = "Refdenotaalpie"
_ESTILO_TEXTO = "Textonotapie"


def extraer_notas_al_pie(ruta_word):
    """Lee las notas al pie de un Word fuente.

    Devuelve una lista de dicts {ancla, texto}, donde `ancla` es el texto del
    párrafo justo antes de la marca de referencia (para reubicarla en el final)
    y `texto` es la definición de la nota. Lista vacía si no hay notas o falla.
    """
    try:
        doc = Document(ruta_word)
    except Exception:
        return []

    # 1. Mapear id de nota → texto, desde la parte footnotes.
    try:
        fn_part = doc.part.part_related_by(RT.FOOTNOTES)
    except KeyError:
        return []
    try:
        root = etree.fromstring(fn_part.blob)
    except Exception:
        return []

    definiciones = {}
    for fn in root.findall(f"{_w}footnote"):
        if fn.get(f"{_w}type") in _TIPOS_NO_NOTA:
            continue
        fid = fn.get(f"{_w}id")
        texto = "".join(t.text or "" for t in fn.iter(f"{_w}t")).strip()
        if fid is not None and texto:
            definiciones[fid] = texto
    if not definiciones:
        return []

    # 2. Recorrer párrafos: al encontrar una referencia, el texto acumulado de
    #    los runs previos es el anclaje.
    notas = []
    for p in doc.paragraphs:
        acc = ""
        for run in p.runs:
            ref = run._r.find(f".//{_w}footnoteReference")
            if ref is not None:
                fid = ref.get(f"{_w}id")
                if fid in definiciones:
                    ancla = acc.strip()[-50:].strip()
                    if ancla:
                        notas.append({"ancla": ancla, "texto": definiciones[fid]})
            acc += run.text or ""
    return notas


def aplicar_notas_al_pie(doc, notas):
    """Inserta `notas` (lista de {ancla, texto}) como notas al pie reales en `doc`.

    Para cada nota busca el `ancla` en los párrafos del documento y coloca la
    marca de referencia justo después; añade la definición a word/footnotes.xml.
    Devuelve el número de notas insertadas.
    """
    if not notas:
        return 0
    try:
        fn_part = doc.part.part_related_by(RT.FOOTNOTES)
    except KeyError:
        return 0

    root = etree.fromstring(fn_part.blob)
    next_id = 1 + max(
        (int(fn.get(f"{_w}id")) for fn in root.findall(f"{_w}footnote")
         if (fn.get(f"{_w}id") or "").lstrip("-").isdigit()),
        default=0,
    )

    insertadas = 0
    for nota in notas:
        ancla = (nota.get("ancla") or "").strip()
        texto = (nota.get("texto") or "").strip()
        if not ancla or not texto:
            continue
        par, pos = _localizar_ancla(doc, ancla)
        if par is None:
            print(f"[REVISAR] Nota al pie sin ubicación en el Word final "
                  f"(ancla {ancla!r}): {texto[:60]!r}")
            continue
        if not _insertar_ref(par, pos, next_id):
            continue
        root.append(_crear_footnote(next_id, texto))
        next_id += 1
        insertadas += 1

    if insertadas:
        fn_part._blob = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )
    return insertadas


# ── Helpers ───────────────────────────────────────────────────────────────────

def _localizar_ancla(doc, ancla):
    """Devuelve (párrafo, posición_char) donde termina el ancla, o (None, None).

    Tolera diferencias de formato al inicio del ancla probando sufijos cada vez
    más cortos; como último recurso usa el token final (ej. 'RO1')."""
    candidatos = [ancla]
    palabras = ancla.split()
    # sufijos quitando palabras del inicio (el final del ancla es lo distintivo)
    for i in range(1, len(palabras)):
        candidatos.append(" ".join(palabras[i:]))
    for objetivo in candidatos:
        objetivo = objetivo.strip()
        if len(objetivo) < 3:
            continue
        for p in doc.paragraphs:
            idx = p.text.find(objetivo)
            if idx != -1:
                return p, idx + len(objetivo)
    return None, None


def _insertar_ref(par, pos, fid):
    """Inserta la marca de referencia en `pos` (índice de carácter) del párrafo,
    dividiendo el run correspondiente y preservando su formato."""
    acc = 0
    for run in list(par.runs):
        rlen = len(run.text)
        if pos <= acc + rlen:
            offset = pos - acc
            cola = run.text[offset:]
            run.text = run.text[:offset]
            ref_r = _crear_ref_run(fid)
            run._r.addnext(ref_r)
            if cola:
                cola_r = _clonar_run(run._r, cola)
                ref_r.addnext(cola_r)
            return True
        acc += rlen
    # posición al final del párrafo
    if par.runs:
        par.runs[-1]._r.addnext(_crear_ref_run(fid))
        return True
    return False


def _crear_ref_run(fid):
    """<w:r><w:rPr><w:rStyle w:val="Refdenotaalpie"/></w:rPr>
        <w:footnoteReference w:id="fid"/></w:r>"""
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), _ESTILO_REF)
    rpr.append(rstyle)
    r.append(rpr)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(fid))
    r.append(ref)
    return r


def _clonar_run(orig_r, texto):
    """Crea un run con el mismo formato (rPr) que `orig_r` y el texto dado."""
    r = OxmlElement("w:r")
    rpr = orig_r.find(qn("w:rPr"))
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = texto
    r.append(t)
    return r


def _crear_footnote(fid, texto):
    """Construye el elemento <w:footnote> con el estilo de texto de nota al pie."""
    fn = OxmlElement("w:footnote")
    fn.set(qn("w:id"), str(fid))

    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    pstyle = OxmlElement("w:pStyle")
    pstyle.set(qn("w:val"), _ESTILO_TEXTO)
    ppr.append(pstyle)
    p.append(ppr)

    # Run con la marca de número de la nota.
    r_ref = OxmlElement("w:r")
    rpr_ref = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), _ESTILO_REF)
    rpr_ref.append(rstyle)
    r_ref.append(rpr_ref)
    r_ref.append(OxmlElement("w:footnoteRef"))
    p.append(r_ref)

    # Run con el texto (con espacio inicial para separar del número).
    r_txt = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = " " + texto
    r_txt.append(t)
    p.append(r_txt)

    fn.append(p)
    return fn
