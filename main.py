"""Punto de entrada para generar el informe semanal completo."""

import os
from datetime import date, timedelta

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import state
from config import *
from utils.word_utils import *
from utils.excel_utils import *
from core.extractores import *
from core.renderers import *
from utils.excel_utils import _obtener_excel_app, _rangos_tablas_sso_backup_dinamico
from state import _workbooks_abiertos
from core.validador import validar_kpis_vs_excel


# ─────────────────────────────────────────────────────────────────────────────
# Helper compartido: construye el Word desde un dict de informes ya listo.
# ─────────────────────────────────────────────────────────────────────────────
_MSG_PENDIENTE = "[No solicitado. Presumiblemente en espera de envío información]"

# Nombres canónicos de imagen según sección e índice dentro de ella
_TEMP = r"C:\Temp"
_SECCION_GH     = "Gestión Hídrica"
_SECCION_SSO    = "Accidentabilidad"
_SECCION_BACKUP = "Accidentabilidad Back-up"

def _nombre_imagen_canonical(seccion, idx):
    """Dado el nombre de sección y el índice de imagen dentro de ella,
    retorna el nombre de archivo PNG que usa _cache() en _construir_doc."""
    if seccion is None:
        return "tabla_principal.png" if idx == 0 else None
    if seccion == _SECCION_GH:
        return "gestion_hidrica.png" if idx == 0 else None
    if seccion == _SECCION_SSO:
        nombres = ["valor_semanal.png", "valor_mensual.png", "valor_anual.png"]
        return nombres[idx] if idx < len(nombres) else None
    if seccion == _SECCION_BACKUP:
        return f"accidentabilidad_{idx + 1}.png"
    # Es clave de faena (MLP, CEN, ANT, CMZ, FCAB)
    if seccion == "MLP" and idx == 1:
        return "tabla_hidrica_mlp.png"
    return f"tabla_{seccion}.png" if idx == 0 else None


def _extraer_imagenes_a_temp(ruta_word):
    """
    Extrae TODAS las imágenes del Word existente y las guarda en C:\\Temp\\
    con los nombres canónicos usados por _cache() en _construir_doc.
    Las imágenes de secciones seleccionadas serán sobreescritas luego por
    los exportadores de Excel; las no seleccionadas quedan intactas.
    Retorna el número de imágenes extraídas.
    """
    os.makedirs(_TEMP, exist_ok=True)

    # Limpiar archivos canónicos de ejecuciones anteriores para evitar cache estale
    _IMAGENES_CANONICAS = (
        ["tabla_principal.png", "gestion_hidrica.png",
         "valor_semanal.png", "valor_mensual.png", "valor_anual.png",
         "tabla_hidrica_mlp.png"]
        + [f"tabla_{c}.png" for c in CONFIG_COMPANIAS]
        + [f"accidentabilidad_{i}.png" for i in range(1, 20)]
    )
    for _nombre in _IMAGENES_CANONICAS:
        _ruta = os.path.join(_TEMP, _nombre)
        if os.path.exists(_ruta):
            os.remove(_ruta)

    _N2K = {cfg["nombre"]: k for k, cfg in CONFIG_COMPANIAS.items()}
    _ESTILOS_TITULO = {"Título 1 AMSA", "Título 2 AMSA"}

    ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    doc = Document(ruta_word)
    seccion_actual = None   # None = preámbulo (antes de cualquier sección)
    idx_en_seccion = 0
    total = 0

    for para in doc.paragraphs:
        texto = para.text.strip()

        # ── Detectar cambio de sección (solo párrafos con estilo de título) ─
        if texto and para.style.name in _ESTILOS_TITULO:
            # Orden importa: "Accidentabilidad Back-up" antes que "Accidentabilidad"
            nueva_seccion = None
            if texto == _SECCION_BACKUP:
                nueva_seccion = _SECCION_BACKUP
            elif texto == _SECCION_SSO:
                nueva_seccion = _SECCION_SSO
            elif texto == _SECCION_GH:
                nueva_seccion = _SECCION_GH
            elif texto in _N2K:
                nueva_seccion = _N2K[texto]

            if nueva_seccion is not None and nueva_seccion != seccion_actual:
                seccion_actual = nueva_seccion
                idx_en_seccion = 0

        # ── Extraer imágenes de este párrafo ────────────────────────────────
        blips = para._p.findall(f".//{{{ns_a}}}blip")
        for blip in blips:
            rid = blip.get(f"{{{ns_r}}}embed")
            if not rid:
                continue
            try:
                img_part = doc.part.related_parts[rid]
                img_bytes = img_part.blob
            except (KeyError, Exception):
                continue

            nombre = _nombre_imagen_canonical(seccion_actual, idx_en_seccion)
            if nombre:
                ruta_destino = os.path.join(_TEMP, nombre)
                with open(ruta_destino, "wb") as f:
                    f.write(img_bytes)
                idx_en_seccion += 1
                total += 1

    print(f"  → {total} imagen(es) extraída(s) del Word existente a {_TEMP}")
    return total

def _construir_doc(
    informes,            # dict clave→texto_compania (para TODAS las faenas)
    excel_madre,
    excel_indicadores,
    dia_inicio, mes_inicio, dia_fin, mes_fin, year, num_semana,
    carpeta_destino, nombre_final,
    incluir_sso=True,
    incluir_gh=True,
    faenas_con_excel=None,           # None = todas; set/list = solo esas usan Excel para su imagen
    secciones_con_datos_previas=None, # claves que tenían datos reales en el Word previo
):
    """Construye y guarda el documento Word.  No hace preguntas ni resuelve rutas."""
    escribir_fechas_excel(excel_madre, dia_inicio, mes_inicio, dia_fin, mes_fin)

    doc = Document(RUTA_PLANTILLA)
    section = doc.sections[0]
    section.top_margin    = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin   = Cm(1.27)
    section.right_margin  = Cm(1.27)

    p = doc.add_paragraph("Informe Semanal de Operación - Antofagasta PLC", style="Título 1 AMSA")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(24)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x12, 0x6F, 0x7A)

    es_parcial = faenas_con_excel is not None

    # ── Línea de estado: secciones actualizadas vs pendientes ─────────────
    _actualizadas = []
    _pendientes   = []
    _previas = set(secciones_con_datos_previas or [])
    for _clave in ORDEN_OFICIAL:
        _tiene_contenido = bool(informes.get(_clave)) and _MSG_PENDIENTE not in informes.get(_clave, "")
        # Actualizada si: fue seleccionada, O tiene contenido real, O el Word previo la tenía con datos
        if (faenas_con_excel is None and _tiene_contenido) or \
           (faenas_con_excel is not None and (_clave in faenas_con_excel or _tiene_contenido or _clave in _previas)):
            _actualizadas.append(_clave)
        else:
            _pendientes.append(_clave)
    if incluir_sso:
        _actualizadas.append("SSO")
    else:
        _pendientes.append("SSO")
    if incluir_gh:
        _actualizadas.append("Gestión Hídrica")
    else:
        _pendientes.append("Gestión Hídrica")

    if _pendientes:
        p_estado = doc.add_paragraph(style="Normal AMSA")
        p_estado.paragraph_format.line_spacing = 1.0
        p_estado.paragraph_format.space_before = Pt(0)
        p_estado.paragraph_format.space_after  = Pt(10)
        p_estado.paragraph_format.left_indent  = Cm(0)
        p_estado.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if _actualizadas:
            _r = p_estado.add_run(f"Actualizadas: {', '.join(_actualizadas)}")
            _r.font.name  = "Arial"
            _r.font.size  = Pt(10)
            _r.bold       = True
            _r.font.color.rgb = RGBColor(0x12, 0x6F, 0x7A)
            _sep = p_estado.add_run("   |   ")
            _sep.font.name = "Arial"
            _sep.font.size = Pt(10)
        _rp = p_estado.add_run(f"Pendientes: {', '.join(_pendientes)}")
        _rp.font.name  = "Arial"
        _rp.font.size  = Pt(10)
        _rp.bold       = True
        _rp.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

    resumen_texto = extraer_resumen_excel(excel_madre)
    for linea in resumen_texto.split("\n"):
        linea_limpia = linea.strip()
        if linea_limpia:
            agregar_texto(doc, linea_limpia)
            if linea_limpia.endswith("."):
                doc.add_paragraph()

    for _ in range(3):
        doc.add_paragraph()

    img_resumen = exportar_imagen_excel(excel_madre, "Grupo Minero FCAB PLAN", "A3:X34", "tabla_principal.png")
    agregar_imagen(doc, img_resumen, 19, 8.8, "")

    if INCLUIR_ESTADO_FASES_DESARROLLO:
        agregar_estado_fases_desarrollo(doc, excel_madre)

    doc.add_page_break()
    agregar_titulo(doc, "Gestión Hídrica", nivel=2)
    if incluir_gh:
        img_hidrica = exportar_imagen_excel(excel_madre, "Gestión Hídrica", "A3:W20", "gestion_hidrica.png")
        agregar_imagen(doc, img_hidrica, 19, 3.24, "")
    else:
        agregar_texto(doc, _MSG_PENDIENTE, color=(128, 128, 128))
        if not es_parcial:
            print("  → Gestión Hídrica: En espera de envío información")

    doc.add_page_break()
    agregar_titulo(doc, "Accidentabilidad", nivel=2)
    if incluir_sso:
        img_semanal = exportar_imagen_excel(excel_indicadores, "Informe Viernes", "A29:M41", "valor_semanal.png")
        img_mensual = exportar_imagen_excel(excel_indicadores, "Informe Viernes", "A15:M27", "valor_mensual.png")
        img_anual   = exportar_imagen_excel(excel_indicadores, "Informe Viernes", "A1:M13",  "valor_anual.png")
        for img_path, texto_titulo in [
            (img_semanal, "Indicadores Valor Semanal"),
            (img_mensual, "Indicadores Valor Mensual"),
            (img_anual,   "Indicadores Valor Anual"),
        ]:
            p_titulo = doc.add_paragraph()
            p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p_titulo.add_run(texto_titulo)
            run.bold = False
            run.font.name = "Arial"
            run.font.size = Pt(11)
            doc.add_paragraph()
            agregar_imagen(doc, img_path, 19, 4.3)
            doc.add_paragraph()
    elif es_parcial:
        # Modo Word existente sin SSO seleccionada: restaurar imágenes del Word previo
        for img_name, texto_titulo in [
            ("valor_semanal.png", "Indicadores Valor Semanal"),
            ("valor_mensual.png", "Indicadores Valor Mensual"),
            ("valor_anual.png",   "Indicadores Valor Anual"),
        ]:
            img_cache = os.path.join(_TEMP, img_name)
            if os.path.exists(img_cache) and os.path.getsize(img_cache) > 0:
                p_titulo = doc.add_paragraph()
                p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p_titulo.add_run(texto_titulo)
                run.bold = False
                run.font.name = "Arial"
                run.font.size = Pt(11)
                doc.add_paragraph()
                agregar_imagen(doc, img_cache, 19, 4.3)
                doc.add_paragraph()
        if not any(
            os.path.exists(os.path.join(_TEMP, n))
            for n in ("valor_semanal.png", "valor_mensual.png", "valor_anual.png")
        ):
            agregar_texto(doc, _MSG_PENDIENTE, color=(128, 128, 128))
    else:
        agregar_texto(doc, _MSG_PENDIENTE, color=(128, 128, 128))
        print("  → Accidentabilidad (SSO): En espera de envío información")

    for clave in ORDEN_OFICIAL:
        cfg = CONFIG_COMPANIAS[clave]
        texto_compania = informes.get(clave, "")
        doc.add_page_break()
        agregar_titulo(doc, cfg["nombre"], nivel=1)
        es_seleccionada = (faenas_con_excel is None or clave in faenas_con_excel)
        if not texto_compania or (not es_seleccionada and _MSG_PENDIENTE in texto_compania):
            agregar_texto(doc, _MSG_PENDIENTE, color=(128, 128, 128))
            continue
        procesador = PROCESADORES_FAENA.get(clave)
        if procesador:
            # Solo regenerar imagen Excel para las faenas seleccionadas
            excel_para_faena = excel_madre if (faenas_con_excel is None or clave in faenas_con_excel) else None
            procesador(doc, texto_compania, excel_para_faena)

    # Accidentabilidad Back-up — siempre al final, después de todas las faenas
    doc.add_page_break()
    agregar_titulo(doc, "Accidentabilidad Back-up", nivel=1)
    if incluir_sso:
        def _obtener_wb_madre():
            wb = _workbooks_abiertos.get(excel_madre)
            if wb is None:
                wb = _obtener_excel_app().Workbooks.Open(excel_madre, UpdateLinks=0)
                _workbooks_abiertos[excel_madre] = wb
            else:
                # Verificar que el proxy COM sigue activo; si no, re-abrir
                try:
                    _ = wb.Name
                except Exception:
                    del _workbooks_abiertos[excel_madre]
                    wb = _obtener_excel_app().Workbooks.Open(excel_madre, UpdateLinks=0)
                    _workbooks_abiertos[excel_madre] = wb
            return wb

        wb_madre = _obtener_wb_madre()
        rangos_tablas_sso = _rangos_tablas_sso_backup_dinamico(wb_madre.Worksheets("SSO"))
        if not rangos_tablas_sso:
            print("  ! Accidentabilidad Back-up: no se encontraron tablas con datos. Verifica hoja SSO.")
        for i, rango_tabla in enumerate(rangos_tablas_sso):
            nombre_img = f"accidentabilidad_{i + 1}.png"
            wb_madre = _obtener_wb_madre()
            img_backup = exportar_imagen_sso_filtrada(excel_madre, wb_madre.Worksheets("SSO"), rango_tabla, nombre_img)

            # Calcular alto real de la imagen a 19 cm de ancho y limitarlo para que
            # quepa en la página sin desbordarse:
            #   i=0: comparte página con el título → max 24 cm
            #   i>0: página propia               → max 26 cm
            MAX_ALTO = 24.0 if i == 0 else 26.0
            alto_cm = MAX_ALTO
            try:
                from PIL import Image as _PILImg
                with _PILImg.open(img_backup) as _im:
                    _w, _h = _im.size
                if _w > 0:
                    alto_cm = min(19.0 * _h / _w, MAX_ALTO)
            except Exception:
                pass

            if i == 0:
                # Primera imagen: va en la misma página que el título
                agregar_imagen(doc, img_backup, 19, alto_cm, "")
            else:
                # Imágenes siguientes: nueva página usando page_break_before en el
                # propio párrafo de imagen para evitar páginas vacías intermedias
                if os.path.exists(img_backup) and os.path.getsize(img_backup) > 0:
                    p = doc.add_paragraph()
                    p.paragraph_format.page_break_before = True
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p.add_run().add_picture(img_backup, width=Cm(19), height=Cm(alto_cm))

        if rangos_tablas_sso:
            print(f"  ✓ Accidentabilidad Back-up: {len(rangos_tablas_sso)} tabla(s) exportada(s)")
    elif es_parcial:
        # Modo Word existente sin SSO seleccionada: restaurar imágenes del Word previo
        i = 0
        while True:
            img_cache = os.path.join(_TEMP, f"accidentabilidad_{i + 1}.png")
            if not (os.path.exists(img_cache) and os.path.getsize(img_cache) > 0):
                break
            MAX_ALTO = 24.0 if i == 0 else 26.0
            alto_cm = MAX_ALTO
            try:
                from PIL import Image as _PILImg
                with _PILImg.open(img_cache) as _im:
                    _w, _h = _im.size
                if _w > 0:
                    alto_cm = min(19.0 * _h / _w, MAX_ALTO)
            except Exception:
                pass
            if i == 0:
                agregar_imagen(doc, img_cache, 19, alto_cm, "")
            else:
                p = doc.add_paragraph()
                p.paragraph_format.page_break_before = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run().add_picture(img_cache, width=Cm(19), height=Cm(alto_cm))
            i += 1
        if i == 0:
            agregar_texto(doc, _MSG_PENDIENTE, color=(128, 128, 128))
    else:
        agregar_texto(doc, _MSG_PENDIENTE, color=(128, 128, 128))
        print("  → Accidentabilidad Back-up (SSO): En espera de envío información")

    texto_pie = construir_texto_semana(dia_inicio, mes_inicio, dia_fin, mes_fin, year)
    agregar_pie_de_pagina(doc, texto_pie)

    ruta_guardado = os.path.join(carpeta_destino, f"{nombre_final}.docx")
    doc.save(ruta_guardado)
    cerrar_excels()
    print(f"Informe generado en: {ruta_guardado}")


def _construir_dirs_semana_anterior(num_semana, dia_inicio, mes_inicio, dia_fin, mes_fin, year, disco=None):
    """Devuelve informes_dirs de la semana anterior buscando en disco por número de semana."""
    try:
        num_ant = max(1, int(num_semana) - 1)
        rutas_actual = construir_rutas_semana(num_semana, dia_inicio, mes_inicio, dia_fin, mes_fin, year, disco=disco)
        raiz_actual = Path(rutas_actual["raiz"])
        # Buscar en el mismo directorio de mes
        parent = raiz_actual.parent
        prev = [f for f in parent.iterdir()
                if f.is_dir() and f.name.startswith(f"{num_ant}_")] if parent.is_dir() else []
        # Si no está en el mismo mes, recorrer todos los meses del año
        if not prev:
            year_dir = parent.parent
            if year_dir.is_dir():
                for mes_dir in year_dir.iterdir():
                    if mes_dir.is_dir():
                        prev += [f for f in mes_dir.iterdir()
                                 if f.is_dir() and f.name.startswith(f"{num_ant}_")]
        if not prev:
            print(f"  ! Semana anterior ({num_ant}) no encontrada en disco — sin fallback")
            return {}
        raiz_ant = prev[0]
        dirs = {
            "MLP":  raiz_ant / "01 -MLP",
            "CEN":  raiz_ant / "02 -CEN",
            "ANT":  raiz_ant / "03 -ANT",
            "CMZ":  raiz_ant / "04 -CMZ",
            "FCAB": raiz_ant / "05 -FCAB",
            "SSO":  raiz_ant / "06 -SSO",
        }
        gh = next((f for f in raiz_ant.iterdir()
                   if f.is_dir() and f.name.startswith("07")), None) if raiz_ant.is_dir() else None
        if gh:
            dirs["Gestión Hídrica"] = gh
        return dirs
    except Exception as e:
        print(f"  ! No se pudo construir fallback de semana anterior: {e}")
        return {}


# ─────────────────────────────────────────────────────────────────────────────
# Actualizar secciones de un Word existente (modo parcial)
# ─────────────────────────────────────────────────────────────────────────────
def actualizar_secciones_word(
    ruta_existente,
    faenas_actualizar,
    dia_inicio, mes_inicio, dia_fin, mes_fin, year, num_semana,
    excel_madre, excel_indicadores, carpeta_destino,
    nombre_override=None, actualizar_vinculos=False,
    informes_paths=None,        # dict clave → ruta del Word fuente de cada faena a actualizar
    excels_dirs_override=None,  # dict clave → directorio/archivo Excel para sobrescribir auto-detección
    incluir_sso=True,
    incluir_gh=True,
    disco=None,
):
    """
    Carga un Word existente y regenera solo las secciones de las faenas indicadas.
    Las secciones no seleccionadas se extraen del Word existente y se re-procesan.
    Cuando actualizar_vinculos=True se actualizan los vínculos de TODAS las faenas.
    """
    if not Path(ruta_existente).is_file():
        print(f"[ERROR] Word existente no encontrado: {ruta_existente}")
        return

    print(f"\n── Modo Word existente: {Path(ruta_existente).name}")
    _secciones_display = list(faenas_actualizar)
    if incluir_sso:
        _secciones_display.append("SSO")
    if incluir_gh:
        _secciones_display.append("Gestión Hídrica")
    print(f"  Secciones a actualizar: {', '.join(_secciones_display)}")

    # ── Extraer textos del Word existente ─────────────────────────────
    _N2K = {cfg["nombre"]: k for k, cfg in CONFIG_COMPANIAS.items()}
    doc_prev = Document(ruta_existente)
    informes_previos, clave_actual, buf = {}, None, []

    def _save_buf():
        if clave_actual and buf:
            informes_previos[clave_actual] = "\n".join(buf)

    # Encabezados de secciones no-faena que marcan el fin del bloque de una compañía
    _SECCIONES_STOP = {_SECCION_GH, _SECCION_SSO, _SECCION_BACKUP}

    for p in doc_prev.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if t in _N2K:
            # Siempre detectar: los nombres de compañías son únicos en el documento
            _save_buf()
            clave_actual = _N2K[t]
            buf = []
        elif t in _SECCIONES_STOP and clave_actual is None:
            # Solo detener antes de entrar a una compañía (nivel documento).
            # Si clave_actual ya está seteada, "Accidentabilidad" es un subtítulo
            # de Hechos Relevantes → se agrega al buf como contenido normal.
            pass
        elif clave_actual:
            buf.append(t)
    _save_buf()

    con_datos = [k for k, v in informes_previos.items() if _MSG_PENDIENTE not in v]
    if con_datos:
        print(f"  Secciones detectadas en Word previo: {', '.join(con_datos)}")
    elif informes_previos:
        print(f"  ⚠ Word previo detectado pero todas las secciones están pendientes")
    else:
        print(f"  ⚠ No se detectaron secciones de compañía en el Word existente")

    # ── Construir dict de informes: fresh para las elegidas, previo para el resto ─
    informes = {}
    for clave in ORDEN_OFICIAL:
        if clave in faenas_actualizar:
            ruta_src = (informes_paths or {}).get(clave, "")
            if ruta_src and Path(ruta_src).is_file():
                informes[clave] = extraer_texto_word(ruta_src)
                print(f"  ✓ {clave}: texto extraído de Word fuente")
            else:
                texto_prev = informes_previos.get(clave, "")
                informes[clave] = texto_prev
                if texto_prev:
                    print(f"  ⚠ {clave}: sin Word fuente → reutilizando texto del Word existente")
                else:
                    print(f"  ⚠ {clave}: sin Word fuente y sin sección previa → quedará vacío")
        else:
            informes[clave] = informes_previos.get(clave, "")
            if informes[clave]:
                print(f"  → {clave}: mantenido del Word existente")

    # ── Actualizar vínculos Excel — TODAS las faenas + SSO + GH ──────────
    if actualizar_vinculos:
        rutas = construir_rutas_semana(num_semana, dia_inicio, mes_inicio, dia_fin, mes_fin, year, disco=disco)
        # Empezar con los dirs de las 5 faenas
        dirs_vinculos = dict(rutas["informes_dirs"])
        # Incluir siempre SSO y Gestión Hídrica para actualizar sus vínculos
        # (aunque no se hayan seleccionado para regenerar imágenes)
        dirs_vinculos.update(rutas.get("excels_adicionales_dirs", {}))
        # Aplicar overrides del usuario: si proporcionó un archivo, tomar su carpeta padre;
        # si proporcionó una carpeta, usarla directamente.
        for clave, ruta_override in (excels_dirs_override or {}).items():
            if not ruta_override:
                continue
            p = Path(ruta_override)
            dirs_vinculos[clave] = str(p.parent) if p.is_file() else str(p)
        if dirs_vinculos:
            # Seleccionadas = faenas elegidas + SSO y GH si se pidió incluirlos
            _seleccionadas = set(faenas_actualizar)
            if incluir_sso:
                _seleccionadas.add("SSO")
            if incluir_gh:
                _seleccionadas.add("Gestión Hídrica")
            abrir_excel_y_actualizar_vinculos(
                excel_madre, dirs_vinculos, carpeta_destino=carpeta_destino,
                ordenar_sso=incluir_sso,
                guardar_en_lugar=True,
                informes_dirs_fallback=_construir_dirs_semana_anterior(
                    num_semana, dia_inicio, mes_inicio, dia_fin, mes_fin, year, disco=disco
                ),
                faenas_seleccionadas=_seleccionadas,
            )

    # ── Extraer imágenes del Word existente → C:\Temp\ con nombres canónicos ─
    # Las secciones NO seleccionadas usarán estas imágenes; las seleccionadas
    # las sobreescribirán con los exportados frescos de Excel.
    _extraer_imagenes_a_temp(ruta_existente)

    nombre_final = nombre_override or os.path.splitext(os.path.basename(ruta_existente))[0]
    _construir_doc(
        informes, excel_madre, excel_indicadores,
        dia_inicio, mes_inicio, dia_fin, mes_fin, year, num_semana,
        carpeta_destino, nombre_final,
        incluir_sso=incluir_sso,
        incluir_gh=incluir_gh,
        faenas_con_excel=set(faenas_actualizar),
        secciones_con_datos_previas=con_datos,
    )


# ─────────────────────────────────────────────────────────────────────────────
# Generar informe semanal completo a partir de los archivos de entrada.
# ─────────────────────────────────────────────────────────────────────────────
def generar_informe(nombre_override=None, incluir_sso=True, incluir_gh=True, disco=None):
    def pedir_entero(mensaje, minimo, maximo):
        while True:
            valor = input(mensaje).strip()
            if valor.isdigit() and minimo <= int(valor) <= maximo:
                return valor
            print(f"  Valor inválido. Ingrese un número entre {minimo} y {maximo}.")

    dia_inicio = pedir_entero("Ingrese el día de inicio: ", 1, 31)
    mes_inicio = pedir_entero("Ingrese el mes de inicio: ", 1, 12)
    dia_fin    = pedir_entero("Ingrese el día de término: ", 1, 31)
    mes_fin    = pedir_entero("Ingrese el mes de término: ", 1, 12)
    year       = pedir_entero("Ingrese el año: ", 2000, 2100)
    num_semana = pedir_entero("Ingrese el número de semana: ", 1, 53)
    texto_pie = f"Semana del {dia_inicio} de {mes_inicio} al {dia_fin} de {mes_fin} {year}"

    orden_oficial = list(ORDEN_OFICIAL)
    seleccion = input(
        "Indica las faenas a procesar separadas por coma (ej: MLP, ANT) o presiona ENTER para todas: "
    ).upper().replace(" ", "")

    if seleccion == "__NINGUNA__":
        faenas_activas = []
    elif seleccion:
        faenas_activas = seleccion.split(",")
    else:
        faenas_activas = orden_oficial

    def _resolver_archivo(ruta_esperada, mensaje):
        """Usa la ruta construida si existe; si no, abre selector."""
        if ruta_esperada and Path(ruta_esperada).is_file():
            print(f"  ✓ {mensaje}: {ruta_esperada.name}")
            return str(ruta_esperada)
        print(f"  ! {mensaje}: no encontrado en ruta esperada → abriendo selector")
        return seleccionar_archivo(mensaje)

    def _resolver_unico_docx(carpeta, mensaje):
        """Busca un único .docx en la carpeta; si hay exactamente uno lo usa, si no abre selector."""
        carpeta = Path(carpeta)
        if carpeta.is_dir():
            docxs = [f for f in carpeta.glob("*.docx") if not f.name.startswith("~$")]
            if len(docxs) == 1:
                print(f"  ✓ {mensaje}: {docxs[0].name}")
                return str(docxs[0])
            if len(docxs) > 1:
                print(f"  ! {mensaje}: {len(docxs)} archivos .docx encontrados → abriendo selector")
        else:
            print(f"  ! {mensaje}: carpeta no encontrada → abriendo selector")
        return seleccionar_archivo(mensaje)

    def _resolver_unico_xlsx(carpeta, mensaje):
        """Busca un .xlsx que empiece con 'BDatos' en la carpeta; si no abre selector."""
        carpeta = Path(carpeta)
        if carpeta.is_dir():
            candidatos = [f for f in carpeta.glob("BDatos*.xlsx") if not f.name.startswith("~$")]
            if len(candidatos) == 1:
                print(f"  ✓ {mensaje}: {candidatos[0].name}")
                return str(candidatos[0])
            if len(candidatos) > 1:
                print(f"  ! {mensaje}: {len(candidatos)} archivos BDatos*.xlsx encontrados → abriendo selector")
            else:
                print(f"  ! {mensaje}: no se encontró BDatos*.xlsx en {carpeta} → abriendo selector")
        else:
            print(f"  ! {mensaje}: carpeta no encontrada → abriendo selector")
        return seleccionar_archivo(mensaje)

    if MODO_DEBUG:
        rutas = construir_rutas_semana(num_semana, dia_inicio, mes_inicio, dia_fin, mes_fin, year, disco=disco)
        excel_madre       = _resolver_archivo(rutas["excel_madre"], "Excel Base")
        excel_indicadores = _resolver_unico_xlsx(rutas["excel_indicadores_dir"], "Excel de indicadores SSO") if incluir_sso else ""
        carpeta_destino   = rutas["carpeta_destino"] if Path(rutas["carpeta_destino"]).is_dir() else seleccionar_carpeta()
        nombre_final      = nombre_override or rutas["nombre_archivo"]
    else:
        excel_madre = seleccionar_archivo("Excel Base")
        excel_indicadores = seleccionar_archivo("Excel de indicadores SSO") if incluir_sso else ""
        carpeta_destino = seleccionar_carpeta()
        nombre_final = nombre_override or input("\nEscribe el nombre del informe final: ")

    if MODO_DEBUG:
        informes_dirs = dict(rutas["informes_dirs"])
        # Incluir SSO y Gestión Hídrica para que sus vínculos también se actualicen
        # y tengan fallback a la semana anterior si no se encuentran en la actual.
        informes_dirs.update(rutas.get("excels_adicionales_dirs", {}))

        nombre_base = os.path.splitext(os.path.basename(excel_madre))[0]
        excel_madre_act = os.path.join(carpeta_destino, f"{nombre_base}_act.xlsx")

        actualizar = input("\n¿Deseas actualizar los vínculos del Excel? (s/n): ").strip().lower()
        if actualizar == "s":
            _sel_gen = set(faenas_activas)
            if incluir_sso:
                _sel_gen.add("SSO")
            if incluir_gh:
                _sel_gen.add("Gestión Hídrica")
            abrir_excel_y_actualizar_vinculos(
                excel_madre, informes_dirs, carpeta_destino=carpeta_destino,
                ordenar_sso=incluir_sso,
                informes_dirs_fallback=_construir_dirs_semana_anterior(
                    num_semana, dia_inicio, mes_inicio, dia_fin, mes_fin, year, disco=disco
                ),
                faenas_seleccionadas=_sel_gen,
            )

        if os.path.exists(excel_madre_act):
            print(f"  ✓ Copia guardada: {os.path.basename(excel_madre_act)}")
        else:
            print("  ! No se creó copia _act.")
        # Seguimos usando el excel_madre ORIGINAL: su workbook está abierto en
        # _workbooks_abiertos con los vínculos ya actualizados en memoria.
        # Abrir la copia _act con UpdateLinks=0 desde una carpeta diferente
        # rompe las rutas relativas y provoca #REF al guardar.
    else:
        informes_dirs = None

    informes = {}
    for clave in orden_oficial:
        if clave in faenas_activas:
            if MODO_DEBUG:
                ruta = _resolver_unico_docx(rutas["informes_dirs"][clave], f"Informe {clave}")
            else:
                ruta = seleccionar_archivo(f"Informe {clave}")
            informes[clave] = extraer_texto_word(ruta) if ruta else ""

    validar = input("\n¿Deseas validar KPIs Word vs Excel? (s/n): ").strip().lower()

    _construir_doc(
        informes, excel_madre, excel_indicadores,
        dia_inicio, mes_inicio, dia_fin, mes_fin, year, num_semana,
        carpeta_destino, nombre_final,
        incluir_sso=incluir_sso,
        incluir_gh=incluir_gh,
    )

    if validar == "s":
        wb_madre = _workbooks_abiertos.get(excel_madre)
        if wb_madre is None:
            wb_madre = _obtener_excel_app().Workbooks.Open(excel_madre, UpdateLinks=0)
            _workbooks_abiertos[excel_madre] = wb_madre
        validar_kpis_vs_excel(informes, wb_madre)

if __name__ == "__main__":
    generar_informe()
