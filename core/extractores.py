"""Funciones de extracción de texto desde Word y desde bloques internos del informe."""

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import re

from state import errores
from utils.text_utils import normalizar_texto_clave

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# ── Títulos de sección del Word de faena ────────────────────────────────────
# Los títulos no vienen siempre escritos igual: MLP escribió "Planta Desaladora"
# hasta la semana 30 de 2026, "Planta Desaladora:" desde la 31 y "Planta
# desaladora:" en la 34. Comparar el literal dejaba la sección vacía sin avisar
# y, cuando lo que fallaba era el corte de la sección anterior, esa se comía el
# resto del documento (semana 34: Concentradora se llevó desaladora e hídrica).
# Por eso la comparación se hace sobre el texto normalizado —sin marcador de
# viñeta, sin tildes, en minúsculas y sin los dos puntos del final—.
def _clave_titulo(linea):
    return normalizar_texto_clave(linea or "").rstrip(":").strip()

# Indica si la línea es exactamente uno de esos títulos de sección.
def es_titulo_seccion(linea, *titulos):
    clave = _clave_titulo(linea)
    return bool(clave) and any(clave == _clave_titulo(t) for t in titulos)

# Indica si la línea abre con uno de esos títulos: el título solo o el título
# seguido de su contenido en la misma línea.
def empieza_titulo_seccion(linea, *titulos):
    clave = _clave_titulo(linea)
    return bool(clave) and any(clave.startswith(_clave_titulo(t)) for t in titulos)

# Avisa de una tabla cuyo texto no entra al informe, con su primera celda para
# poder ubicarla en el Word de faena.
def _avisar_tabla_omitida(tabla, columnas):
    primera = next((c.text.strip() for f in tabla.rows for c in f.cells if c.text.strip()), "")
    msg = (f"[REVISAR] Word de faena: tabla de {len(tabla.rows)}x{columnas} omitida al "
           f"extraer el texto -> '{primera[:80]}'")
    print(msg)
    errores.append(msg)

# Recorre el cuerpo del Word en orden de documento y devuelve sus párrafos,
# incluidos los que van dentro de una tabla de una sola columna: el redactor a
# veces encierra un bloque de texto en una tabla así para encuadrarlo —MLP metió
# toda la Concentradora de la semana 34 de 2026 en una— y doc.paragraphs no los
# ve, con lo que la sección entera desaparecía del informe sin dejar rastro. Las
# tablas de más de una columna sí son datos tabulares, que el informe toma del
# Excel madre y no del Word: se omiten con aviso, para no meter celdas sueltas
# como viñetas.
def parrafos_en_orden(doc, avisar_tablas=False):
    for hijo in doc.element.body.iterchildren():
        if hijo.tag == f"{_W}p":
            yield Paragraph(hijo, doc)
        elif hijo.tag == f"{_W}tbl":
            tabla = Table(hijo, doc)
            try:
                columnas = len(tabla.columns)
                filas = list(tabla.rows)
            except Exception:
                continue
            if columnas != 1:
                if avisar_tablas:
                    _avisar_tabla_omitida(tabla, columnas)
                continue
            for fila in filas:
                for celda in fila.cells:
                    for p in celda.paragraphs:
                        yield p

# Lee un documento Word y devuelve su texto completo.
def extraer_texto_word(ruta_word):
    try:
        doc = Document(ruta_word)
        lineas = [p.text.strip() for p in parrafos_en_orden(doc, avisar_tablas=True)]
        return "\n".join(l for l in lineas if l)
    except Exception as e:
        errores.append(f"[ERROR] No se pudo leer el informe {ruta_word}: {e}")
        return ""

# Extrae un bloque de texto entre un título inicial y uno final.
def extraer_bloque(texto, inicio, finales=()):
    seccion = []
    capturar = False
    for linea in texto.split("\n"):
        l = linea.strip()
        if not capturar:
            if inicio in l:
                capturar = True
                continue
        else:
            if finales and any(l.startswith(f) or l == f for f in finales):
                break
            if l:
                seccion.append(l)
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_accidentabilidad(texto):
    return extraer_bloque(texto, inicio="Accidentabilidad", finales=("Reportabilidad",))

# Extrae información específica desde el texto o archivo de origen.
def extraer_reportabilidad(texto):
    # "Medio Ambiente:" (con ":" inmediato) es sub-ítem de Reportabilidad → no corta.
    # "Medio Ambiente" sin ":" es encabezado standalone → sí corta.
    FINALES_PREFIJO = ("Gestión SSO", "Salud Ocupacional y Gestión Vial",
                       "Producción Semana", "Asuntos Públicos")
    seccion = []
    capturar = False
    for linea in texto.split("\n"):
        l = linea.strip()
        if not capturar:
            if "Reportabilidad" in l:
                capturar = True
                continue
        else:
            if l.startswith("Medio Ambiente"):
                break
            if any(l.startswith(f) for f in FINALES_PREFIJO):
                break
            if l:
                seccion.append(l)
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_medio_ambiente(texto):
    # Solo activar en encabezado standalone "Medio Ambiente" (sin ":" inmediato).
    # "Medio Ambiente:" como sub-ítem dentro de Reportabilidad NO dispara esto.
    seccion = []
    capturar = False
    finales = ("Asuntos Públicos", "Gestión SSO", "Producción Semana")
    for linea in texto.split("\n"):
        l = linea.strip()
        if not capturar:
            if l.startswith("Medio Ambiente"):
                capturar = True
                continue
        else:
            if any(l.startswith(f) for f in finales):
                break
            if l:
                l_limpia = re.sub(r"^[•\-\·\s]*", "", l)
                seccion.append(l_limpia)
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_asuntos_publicos(texto):
    return extraer_bloque(texto, inicio="Asuntos Públicos", finales=("Producción Semana",))

# Extrae información específica desde el texto o archivo de origen.
def extraer_gestion_sso(texto):
    return extraer_bloque(
        texto,
        inicio="Gestión SSO",
        finales=(
            "Salud Ocupacional y Gestión Vial",
            "Producción Semana",
            "Medio Ambiente",
        ),
    )

# Extrae información específica desde el texto o archivo de origen.
def extraer_salud_ocupacional(texto):
    return extraer_bloque(texto, inicio="Salud Ocupacional y Gestión Vial", finales=("Medio Ambiente",))

# Extrae información específica desde el texto o archivo de origen.
def extraer_principales_desviaciones(texto):
    seccion = []
    capturar = False
    for linea in texto.split("\n"):
        if linea.startswith("Principales Desviaciones"):
            capturar = True
            continue
        if capturar:
            if linea.startswith("Mina") or linea.startswith("Tren"):
                break
            seccion.append(linea.strip())
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_mina(texto):
    seccion = []
    capturar = False
    en_principales = False
    PARADAS_EXACTAS = {"Planta", "Planta:"}
    PARADAS_PREFIJO = ("Concentradora", "Sulfuros", "Detalle por fases")
    for linea in texto.split("\n"):
        l = linea.strip()
        if l.startswith("Principales Desviaciones"):
            en_principales = True
            continue
        if not capturar:
            if linea.startswith("Mina"):
                # Solo disparar si es encabezado corto, no contenido ("Mina: texto largo...")
                resto = l[4:].lstrip(":").strip()
                if len(resto) < 10:
                    capturar = True
                    continue
            # Sin encabezado: fallback en "Movimiento Mina" solo dentro de PD
            if en_principales and l.startswith("Movimiento Mina"):
                capturar = True
        if capturar:
            if l in PARADAS_EXACTAS or any(l.startswith(p) for p in PARADAS_PREFIJO):
                break
            seccion.append(l)
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_concentradora(texto):
    seccion = []
    capturar = False
    for linea in texto.split("\n"):
        l = linea.strip()
        if not capturar:
            if es_titulo_seccion(l, "Concentradora"):
                capturar = True
            continue
        # "Gestión Hídrica" también corta, por si esa semana no viene desaladora.
        if empieza_titulo_seccion(l, "Planta Desaladora", "Gestión Hídrica"):
            break
        seccion.append(l)
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_sulfuros(texto):
    seccion = []
    capturar = False
    for linea in texto.split("\n"):
        if linea.startswith("Sulfuros"):
            capturar = True
            continue
        if capturar:
            if linea.startswith("Cátodos"):
                break
            seccion.append(linea.strip())
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_cátodos(texto):
    seccion = []
    capturar = False
    for linea in texto.split("\n"):
        if linea.startswith("Cátodos"):
            capturar = True
            continue
        if capturar:
            seccion.append(linea.strip())
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_detalle_fases(texto):
    seccion = []
    capturar = False
    for linea in texto.split("\n"):
        if linea.startswith("Detalle por fases"):
            capturar = True
            continue
        if capturar:
            if linea.startswith("Planta") or linea.startswith("Planta:"):
                break
            seccion.append(linea.strip())
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_planta(texto):
    seccion = []
    capturar = False
    en_principales = False
    for linea in texto.split("\n"):
        l = linea.strip()
        if l.startswith("Principales Desviaciones"):
            en_principales = True
            continue
        if not en_principales:
            continue
        if l.startswith("Planta:") or (l.startswith("Planta") and not l.startswith("Planta Desaladora") and not l.startswith("Planta Hidro")):
            capturar = True
            continue
        if capturar:
            seccion.append(l)
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_planta_desaladora(texto):
    seccion = []
    capturar = False
    for linea in texto.split("\n"):
        l = linea.strip()
        if not capturar:
            # Solo el título suelto abre la sección: "En Planta Desaladora el
            # flujo..." es contenido y no debe disparar la captura.
            if es_titulo_seccion(l, "Planta Desaladora"):
                capturar = True
            continue
        if empieza_titulo_seccion(l, "Gestión Hídrica"):
            break
        seccion.append(l)
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_gestión_hídrica(texto):
    seccion = []
    capturar = False
    for linea in texto.split("\n"):
        l = linea.strip()
        if not capturar:
            if es_titulo_seccion(l, "Gestión Hídrica"):
                capturar = True
            continue
        seccion.append(l)
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_tren(texto):
    seccion = []
    capturar = False
    for linea in texto.split("\n"):
        if linea.strip() == "Tren" or linea.strip() == "Tren:":
            capturar = True
            continue
        if capturar:
            if linea.startswith("Camión") or linea.startswith("Camión:"):
                break
            seccion.append(linea.strip())
    return seccion

# Extrae información específica desde el texto o archivo de origen.
def extraer_camión(texto):
    seccion = []
    capturar = False
    for linea in texto.split("\n"):
        if linea.strip() == "Camión" or linea.strip() == "Camión:":
            capturar = True
            continue
        if capturar:
            seccion.append(linea.strip())
    return seccion
