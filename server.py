"""
Panel de Control Web — Informe Semanal Antofagasta Minerals
Uso: python server.py
     Luego abrir http://localhost:5000 en el navegador.

Dependencias adicionales (instalar una vez):
    pip install flask pymupdf pyspellchecker
"""

import builtins
import importlib
import io
import os
import sys
import threading
import traceback
from pathlib import Path

from flask import Flask, jsonify, request, send_from_directory

BASE_DIR = Path(__file__).parent
app = Flask(__name__, static_folder=str(BASE_DIR))

# ── Log buffer ────────────────────────────────────────────────────────────────
_logs: list[str] = []
_lock   = threading.Lock()
_running = False
_orig_stdout = sys.stdout
_orig_input  = builtins.input

class _Tee(io.TextIOBase):
    def write(self, s):
        _orig_stdout.write(s); _orig_stdout.flush()
        if s:
            with _lock: _logs.append(s)
        return len(s)
    def flush(self): _orig_stdout.flush()

def _start_cap(): sys.stdout = _Tee()
def _stop_cap():  sys.stdout = _orig_stdout

# ── Rutas ─────────────────────────────────────────────────────────────────────
@app.route("/")
def index():
    return send_from_directory(str(BASE_DIR), "panel.html")

# ── Logs ──────────────────────────────────────────────────────────────────────
@app.route("/api/logs")
def api_logs():
    offset = int(request.args.get("offset", 0))
    with _lock: chunk = _logs[offset:]
    return jsonify({"lines": chunk, "next_offset": offset + len(chunk), "running": _running})

@app.route("/api/logs/clear", methods=["POST"])
def api_clear():
    with _lock: _logs.clear()
    return jsonify({"ok": True})

# ── Browse (abre dialogo nativo via subprocess para evitar problemas de hilos) ─
@app.route("/api/browse-file", methods=["POST"])
def api_browse_file():
    d     = request.json
    title = d.get("title", "Seleccionar archivo")
    tipos = d.get("tipos", "todos")
    ftype_map = {
        "excel": "[('Excel','*.xlsx *.xlsm'),('Todos','*.*')]",
        "word":  "[('Word','*.docx'),('Todos','*.*')]",
        "pdf":   "[('PDF','*.pdf'),('Todos','*.*')]",
        "todos": "[('Todos','*.*')]",
    }
    ft = ftype_map.get(tipos, ftype_map["todos"])
    script = (
        "import tkinter as tk; from tkinter import filedialog; "
        "root=tk.Tk(); root.withdraw(); root.wm_attributes('-topmost',1); "
        f"f=filedialog.askopenfilename(title={title!r},filetypes={ft}); "
        "print(f,end='')"
    )
    import subprocess
    r = subprocess.run([sys.executable, "-c", script], capture_output=True, text=True)
    return jsonify({"path": r.stdout.strip()})

@app.route("/api/browse-folder", methods=["POST"])
def api_browse_folder():
    title = request.json.get("title", "Seleccionar carpeta")
    script = (
        "import tkinter as tk; from tkinter import filedialog; "
        "root=tk.Tk(); root.withdraw(); root.wm_attributes('-topmost',1); "
        f"f=filedialog.askdirectory(title={title!r}); "
        "print(f,end='')"
    )
    import subprocess
    r = subprocess.run([sys.executable, "-c", script], capture_output=True, text=True)
    return jsonify({"path": r.stdout.strip()})

# ── Info de semana (devuelve todas las rutas auto-detectadas) ─────────────────
@app.route("/api/semana-info", methods=["POST"])
def api_semana_info():
    d = request.json
    try:
        from config import construir_rutas_semana
        from utils.excel_utils import _PATRON_EXCEL_FAENA

        disco = d.get("disco") or None
        rutas = construir_rutas_semana(
            d["num_semana"], d["dia_inicio"], d["mes_inicio"],
            d["dia_fin"],    d["mes_fin"],    d["year"],
            disco=disco,
        )

        # Si el usuario eligió la carpeta manualmente, usarla como raíz
        raiz_override = (d.get("raiz_override") or "").strip()
        if raiz_override:
            raiz_base = Path(raiz_override)
            # Reconstruir todas las sub-rutas relativas dentro de la raíz elegida
            rutas["raiz"]                = raiz_base
            rutas["excel_madre"]         = raiz_base / Path(rutas["excel_madre"]).name
            rutas["excel_indicadores_dir"] = raiz_base / "06 -SSO"
            rutas["carpeta_destino"]     = str(raiz_base)
            rutas["informes_dirs"]       = {
                "MLP":  raiz_base / "01 -MLP",
                "CEN":  raiz_base / "02 -CEN",
                "ANT":  raiz_base / "03 -ANT",
                "CMZ":  raiz_base / "04 -CMZ",
                "FCAB": raiz_base / "05 -FCAB",
            }
            # Gestión Hídrica: primera subcarpeta que empiece con "07"
            gh_dir = next((f for f in raiz_base.iterdir()
                           if f.is_dir() and f.name.startswith("07")), None) if raiz_base.is_dir() else None
            rutas["excels_adicionales_dirs"] = {
                "SSO":             raiz_base / "06 -SSO",
                "Gestión Hídrica": gh_dir or raiz_base / "07 -Gestión Hídrica",
            }

        raiz        = Path(rutas["raiz"])
        excel_madre = Path(rutas["excel_madre"])
        ind_dir     = Path(rutas["excel_indicadores_dir"])

        # Excel indicadores: buscar BDatos*.xlsx
        cands_ind = list(ind_dir.glob("BDatos*.xlsx")) if ind_dir.is_dir() else []
        excel_ind = str(cands_ind[0]) if len(cands_ind) == 1 else ""

        # Informes Word por faena
        informes     = {}
        excels_faena = {}

        for clave, sub in rutas["informes_dirs"].items():
            sub = Path(sub)
            # Word
            docxs = [f for f in sub.glob("*.docx") if not f.name.startswith("~$")] if sub.is_dir() else []
            informes[clave] = str(docxs[0]) if len(docxs) == 1 else ""
            # Excel vinculado (faenas principales)
            patron = _PATRON_EXCEL_FAENA.get(clave, "").lower()
            if patron and sub.is_dir():
                cands = [f for f in sub.glob("*.xlsx")
                         if not f.name.startswith("~$") and patron in f.name.lower()]
                excels_faena[clave] = str(cands[0]) if len(cands) == 1 else ""
            else:
                excels_faena[clave] = ""

        # Excels adicionales: SSO y Gestión Hídrica
        for clave, adir in rutas.get("excels_adicionales_dirs", {}).items():
            adir = Path(adir)
            patron = _PATRON_EXCEL_FAENA.get(clave, "").lower()
            if patron and adir.is_dir():
                cands = [f for f in adir.glob("*.xlsx")
                         if not f.name.startswith("~$") and patron in f.name.lower()]
                excels_faena[clave] = str(cands[0]) if len(cands) == 1 else ""
            else:
                excels_faena[clave] = ""

        return jsonify({
            "raiz":                str(raiz),
            "raiz_ok":             raiz.is_dir(),
            "excel_madre":         str(excel_madre),
            "excel_madre_ok":      excel_madre.is_file(),
            "excel_indicadores":   excel_ind,
            "excel_indicadores_ok": bool(excel_ind),
            "carpeta_destino":     str(raiz),
            "carpeta_destino_ok":  raiz.is_dir(),
            "nombre_archivo":      rutas["nombre_archivo"],
            "informes":            informes,
            "excels_faena":        excels_faena,
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 400

# ── Calendario de semanas ─────────────────────────────────────────────────────
@app.route("/api/calendario")
def api_calendario():
    """Lee Calendario Informe Semanal.xlsx y devuelve fechas por número de semana."""
    cal_path = BASE_DIR / "Calendario Informe Semanal.xlsx"
    if not cal_path.is_file():
        return jsonify({"error": "no_encontrado", "semanas": {}})
    try:
        import openpyxl
        import re as _re
        wb = openpyxl.load_workbook(str(cal_path), read_only=True, data_only=True)
        semanas = {}
        COL_SEMANA  = 17   # columna R (0-based)
        COL_DETALLE = 18   # columna S (0-based)
        pat_ini = _re.compile(r'[Rr]eales del\s+(\d{1,2})/(\d{1,2})')
        pat_fin = _re.compile(r'[Pp]royectado\s+(\d{1,2})/(\d{1,2})')
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                row = list(row)
                if len(row) <= COL_DETALLE:
                    continue
                num_raw = row[COL_SEMANA]
                det_raw = row[COL_DETALLE]
                if num_raw is None or det_raw is None:
                    continue
                try:
                    num = int(float(num_raw))
                except (TypeError, ValueError):
                    continue
                if not (1 <= num <= 53):
                    continue
                det_str = str(det_raw).strip()
                m_ini = pat_ini.search(det_str)
                m_fin = pat_fin.search(det_str)
                if m_ini and m_fin and num not in semanas:
                    semanas[num] = {
                        "dia_inicio": int(m_ini.group(1)),
                        "mes_inicio": int(m_ini.group(2)),
                        "dia_fin":    int(m_fin.group(1)),
                        "mes_fin":    int(m_fin.group(2)),
                        "detalles":   det_str,
                    }
        wb.close()
        return jsonify({"semanas": semanas})
    except Exception as e:
        return jsonify({"error": str(e), "semanas": {}})

# ── Generar ───────────────────────────────────────────────────────────────────
@app.route("/api/generar", methods=["POST"])
def api_generar():
    global _running
    if _running: return jsonify({"error": "Ya hay una tarea en ejecución."}), 409
    t = threading.Thread(target=_task, args=(_generar, request.json), daemon=True)
    t.start()
    return jsonify({"ok": True})

# ── Validar ───────────────────────────────────────────────────────────────────
@app.route("/api/validar", methods=["POST"])
def api_validar():
    global _running
    if _running: return jsonify({"error": "Ya hay una tarea en ejecución."}), 409
    t = threading.Thread(target=_task, args=(_validar, request.json), daemon=True)
    t.start()
    return jsonify({"ok": True})

# ── Resultado validación (panel HTML) ────────────────────────────────────────
@app.route("/api/validar-resultado")
def api_validar_resultado():
    from core.validador import get_resultados
    return jsonify(get_resultados())

# ── Revisar ortografía y gramática ───────────────────────────────────────────
@app.route("/api/revisar-ortografia", methods=["POST"])
def api_revisar_ortografia():
    global _running
    if _running: return jsonify({"error": "Ya hay una tarea en ejecución."}), 409
    t = threading.Thread(target=_task, args=(_revisar_ortografia, request.json), daemon=True)
    t.start()
    return jsonify({"ok": True})

# ── Wrapper de tarea ──────────────────────────────────────────────────────────
def _task(fn, data):
    global _running
    _running = True; _start_cap()
    try:    fn(data)
    except: print(f"\n[ERROR]\n{traceback.format_exc()}")
    finally: _stop_cap(); _running = False

# ── Implementación: Generar ───────────────────────────────────────────────────
_path_overrides: dict = {}

def _patched_sel_archivo(msg=""):
    msg_n = msg.lower()
    for key, path in _path_overrides.items():
        if key and path and key.lower() in msg_n:
            print(f"  ✓ (ruta del panel) {msg}: {Path(path).name}")
            return path
    print(f"  ! No encontrado automáticamente: {msg}")
    print("    → Puedes proporcionar la ruta en el panel y volver a generar.")
    return None

def _patched_sel_carpeta():
    return _path_overrides.get("__destino__", "")

def _mock_input(responses: list):
    resp = list(responses)
    def _inp(prompt=""):
        val = resp.pop(0) if resp else ""
        return str(val)
    return _inp

def _generar(d):
    global _path_overrides

    from config import construir_rutas_semana
    rutas = construir_rutas_semana(
        d["num_semana"], d["dia_inicio"], d["mes_inicio"],
        d["dia_fin"],    d["mes_fin"],    d["year"],
    )

    carpeta_destino = d.get("carpeta_destino", "").strip() or str(rutas["raiz"])
    word_existente  = d.get("word_existente", "").strip()

    # Construir overrides de rutas desde lo que el usuario llenó en el panel
    _path_overrides = {
        "excel base":    d.get("excel_madre", ""),
        "indicadores":   d.get("excel_indicadores", ""),
        "__destino__":   carpeta_destino,
    }
    for clave, path in d.get("informes", {}).items():
        if path:
            _path_overrides[clave.lower()] = path

    faenas  = d.get("faenas", [])
    act_str = "s" if d.get("actualizar_vinculos") else "n"

    import utils.excel_utils as eu
    _orig_a = eu.seleccionar_archivo
    _orig_c = eu.seleccionar_carpeta

    eu.seleccionar_archivo = _patched_sel_archivo
    eu.seleccionar_carpeta = _patched_sel_carpeta

    try:
        import main as m
        importlib.reload(m)
        nombre_custom = d.get("nombre_archivo", "").strip() or None

        if word_existente:
            # ── Modo Word existente: actualizar secciones seleccionadas ──
            if not Path(word_existente).is_file():
                print(f"[ERROR] Word existente no encontrado: {word_existente}")
                return

            excel_madre     = d.get("excel_madre", "").strip()
            excel_ind       = d.get("excel_indicadores", "").strip()
            informes_paths  = {k: v for k, v in d.get("informes", {}).items() if v}

            if not excel_madre or not Path(excel_madre).is_file():
                print(f"[ERROR] Excel madre no encontrado: {excel_madre}")
                return
            if d.get("incluir_sso", True) and (not excel_ind or not Path(excel_ind).is_file()):
                print(f"[ERROR] Excel indicadores no encontrado: {excel_ind}")
                return
            if not faenas:
                print("[ERROR] No se seleccionaron secciones a actualizar.")
                return

            m.actualizar_secciones_word(
                ruta_existente      = word_existente,
                faenas_actualizar   = faenas,
                dia_inicio          = d["dia_inicio"],
                mes_inicio          = d["mes_inicio"],
                dia_fin             = d["dia_fin"],
                mes_fin             = d["mes_fin"],
                year                = d["year"],
                num_semana          = d["num_semana"],
                excel_madre         = excel_madre,
                excel_indicadores   = excel_ind,
                carpeta_destino     = carpeta_destino,
                nombre_override     = nombre_custom,
                actualizar_vinculos = d.get("actualizar_vinculos", False),
                informes_paths      = informes_paths,
                incluir_sso         = d.get("incluir_sso", True),
                incluir_gh          = d.get("incluir_gh",  True),
            )
        else:
            # ── Modo Word nuevo: generación completa (comportamiento original) ──
            if not Path(rutas["raiz"]).is_dir():
                print(f"[ERROR] Carpeta de semana no encontrada:\n  {rutas['raiz']}")
                print("  Verifica la ruta base y los datos de semana ingresados.")
                return

            respuestas = [
                str(d["dia_inicio"]), str(d["mes_inicio"]),
                str(d["dia_fin"]),    str(d["mes_fin"]),
                str(d["year"]),       str(d["num_semana"]),
                ",".join(faenas) if faenas else "",
                act_str,
                "n",   # validar KPIs — siempre NO, es módulo separado
            ]
            builtins.input = _mock_input(respuestas)
            m.generar_informe(
                nombre_override = nombre_custom,
                incluir_sso     = d.get("incluir_sso", True),
                incluir_gh      = d.get("incluir_gh",  True),
            )

        # Verificar formato del documento generado
        nombre_final = nombre_custom or (
            os.path.splitext(os.path.basename(word_existente))[0]
            if word_existente else rutas["nombre_archivo"]
        )
        ruta_docx = os.path.join(carpeta_destino, f"{nombre_final}.docx")
        if Path(ruta_docx).is_file():
            _verificar_documento(ruta_docx)

    finally:
        builtins.input         = _orig_input
        eu.seleccionar_archivo = _orig_a
        eu.seleccionar_carpeta = _orig_c
        _path_overrides        = {}

# ── Verificación de formato del documento generado ────────────────────────────
def _verificar_documento(ruta_docx):
    """Verifica convenciones de formato en el Word recién generado."""
    import re
    from docx import Document

    print("\n── Verificación del documento ──────────────────────────────")
    try:
        doc = Document(ruta_docx)
    except Exception as e:
        print(f"  ! No se pudo abrir para verificación: {e}")
        return

    texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # 1) Coma como separador decimal (ej: "3,5" "12,4%")
    #    Se excluyen separadores de miles (\d{3} tras la coma)
    coma_dec = re.findall(r'\b\d+,\d{1,2}\b(?!\d)', texto)
    if coma_dec:
        muestra = ", ".join(dict.fromkeys(coma_dec)[:5])
        print(f"  [REVISAR] Coma usada como decimal (usar punto): {muestra}")
    else:
        print("  ✓ Decimales: no se detectaron comas como separador decimal")

    # 2) Espacios dobles
    dobles = re.findall(r'  +', texto)
    if dobles:
        print(f"  [REVISAR] {len(dobles)} ocurrencia(s) de espacios dobles en el texto")
    else:
        print("  ✓ Espaciado: sin espacios dobles detectados")

    # 3) Porcentajes sin espacio antes del símbolo (ej: "5%" vs "5 %")
    #    Convención: sin espacio → si hay espacio, avisar
    espacio_pct = re.findall(r'\d\s+%', texto)
    if espacio_pct:
        print(f"  [REVISAR] {len(espacio_pct)} caso(s) con espacio antes de '%' (ej: '5 %' → '5%')")
    else:
        print("  ✓ Porcentajes: formato correcto (sin espacio antes de '%')")

    print("  ✓ Verificación de formato completada")


# ── Implementación: Validar ───────────────────────────────────────────────────
def _validar(d):
    ruta_word  = d.get("ruta_word",  "").strip()
    ruta_excel = d.get("ruta_excel", "").strip()
    faenas     = d.get("faenas") or None   # None → valida todo

    if not Path(ruta_word).is_file():
        print(f"[ERROR] Word no encontrado: {ruta_word}"); return
    if not Path(ruta_excel).is_file():
        print(f"[ERROR] Excel no encontrado: {ruta_excel}"); return

    from docx import Document
    from config import CONFIG_COMPANIAS
    from core.validador import validar_kpis_vs_excel
    from utils.excel_utils import _obtener_excel_app
    from state import _workbooks_abiertos

    print(f"\n── Validación KPIs ────────────────────────────────────────")
    print(f"  Word : {Path(ruta_word).name}")
    print(f"  Excel: {Path(ruta_excel).name}")
    if faenas:
        print(f"  Secciones: {', '.join(faenas)}")
    print()

    _N2K = {cfg["nombre"]: k for k, cfg in CONFIG_COMPANIAS.items()}
    doc = Document(ruta_word)
    informes, clave_actual, buf = {}, None, []

    def _save():
        if clave_actual and buf:
            informes[clave_actual] = "\n".join(buf)

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        if t in _N2K:
            _save(); clave_actual = _N2K[t]; buf = []
        elif clave_actual:
            buf.append(t)
    _save()

    # Filtrar solo las faenas solicitadas
    if faenas:
        informes = {k: v for k, v in informes.items() if k in faenas}

    if not informes:
        print("  ! No se detectaron secciones de compañía en el Word."); return
    print(f"  Compañías a validar: {', '.join(informes)}")

    ruta_abs = str(Path(ruta_excel).resolve())
    wb = _workbooks_abiertos.get(ruta_abs)
    if wb is None:
        print("  Abriendo Excel...")
        wb = _obtener_excel_app().Workbooks.Open(ruta_abs, UpdateLinks=0)
        _workbooks_abiertos[ruta_abs] = wb

    validar_kpis_vs_excel(informes, wb)

# ── Implementación: Revisar ortografía y gramática ───────────────────────────
def _revisar_ortografia(d):
    ruta   = d.get("ruta_docx", "").strip()
    faenas = d.get("faenas") or None   # None → revisa todo

    if not ruta or not Path(ruta).is_file():
        print(f"[ERROR] Archivo no encontrado: {ruta}"); return

    try:
        from revisar_gramatica import revisar_gramatica
    except ImportError as e:
        print(f"[ERROR] No se pudo importar revisar_gramatica: {e}"); return

    revisar_gramatica(ruta, faenas=faenas)

# ── Arranque ──────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("\n╔══════════════════════════════════════════════════════╗")
    print("║  Control de Gestión AMSA — Panel Web                ║")
    print("║  Abre en el navegador: http://localhost:5000         ║")
    print("╚══════════════════════════════════════════════════════╝\n")
    app.run(host="127.0.0.1", port=5000, debug=False, threaded=True)
